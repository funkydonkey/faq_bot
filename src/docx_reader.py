"""Utility module for reading and processing DOCX files."""
from docx import Document
from typing import List, Dict


class DocxReader:
    """Class to read and extract content from DOCX files."""

    def __init__(self, file_path: str):
        """Initialize the reader with a DOCX file path."""
        self.file_path = file_path
        self.document = None
        self.content = None

    def load(self):
        """Load the DOCX file."""
        try:
            self.document = Document(self.file_path)
            self._extract_content()
            return True
        except Exception as e:
            print(f"Error loading document: {e}")
            return False

    def _extract_content(self):
        """Extract all text content from the document."""
        paragraphs = []
        for para in self.document.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())

        # Also extract text from tables
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())

        self.content = "\n\n".join(paragraphs)

    def get_content(self) -> str:
        """Get the full text content of the document."""
        return self.content if self.content else ""

    def get_paragraphs(self) -> List[str]:
        """Get content as a list of paragraphs."""
        if not self.content:
            return []
        return [p.strip() for p in self.content.split("\n\n") if p.strip()]

    def search(self, query: str) -> List[Dict[str, str]]:
        """Search for text in the document and return matching paragraphs."""
        if not self.content:
            return []

        results = []
        paragraphs = self.get_paragraphs()
        query_lower = query.lower()

        for idx, para in enumerate(paragraphs):
            if query_lower in para.lower():
                results.append({
                    "paragraph_number": idx + 1,
                    "content": para
                })

        return results
